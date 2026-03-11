# -*- coding: utf-8 -*-
"""
一键显著 - 核心功能模块
支持多种优化策略使回归结果显著：数据层面、模型层面、样本层面、高级方法等
"""

import pandas as pd
import numpy as np
import statsmodels.api as sm
from statsmodels.formula.api import ols
import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
warnings.filterwarnings('ignore')


# ============================================================
# 数据层面优化
# ============================================================
def winsorize_series(series, limits=(0.01, 0.99)):
    """
    Winsorize 缩尾处理
    
    Parameters:
    -----------
    series : Series
        需要缩尾的序列
    limits : tuple
        缩尾的分位点 (下限，上限)
    
    Returns:
    --------
    Series : 缩尾后的序列
    """
    lower = series.quantile(limits[0])
    upper = series.quantile(limits[1])
    return series.clip(lower=lower, upper=upper)


def winsorize_df(df, variables, limits=(0.01, 0.99)):
    """
    对多个变量进行缩尾处理
    
    Parameters:
    -----------
    df : DataFrame
        数据
    variables : list
        需要缩尾的变量列表
    limits : tuple
        缩尾的分位点
    
    Returns:
    --------
    DataFrame : 缩尾后的数据
    """
    df_winsorized = df.copy()
    for var in variables:
        if var in df_winsorized.columns:
            df_winsorized[var] = winsorize_series(df_winsorized[var], limits)
    return df_winsorized


def transform_variable(series, method='log'):
    """
    变量转换
    
    Parameters:
    -----------
    series : Series
        需要转换的序列
    method : str
        转换方法 ('log', 'sqrt', 'reciprocal', 'square')
    
    Returns:
    --------
    Series : 转换后的序列
    """
    if method == 'log':
        # 处理 0 和负值
        if series.min() <= 0:
            return np.log1p(series)  # log(1+x)
        else:
            return np.log(series)
    elif method == 'sqrt':
        return np.sqrt(series.abs())
    elif method == 'reciprocal':
        return 1 / (series.abs() + 1)
    elif method == 'square':
        return series ** 2
    else:
        return series


def standardize_series(series):
    """Z-score 标准化"""
    return (series - series.mean()) / series.std()


def standardize_df(df, variables):
    """对多个变量进行标准化"""
    df_std = df.copy()
    for var in variables:
        if var in df_std.columns:
            df_std[var] = standardize_series(df_std[var])
    return df_std


# ============================================================
# 模型层面优化
# ============================================================
def baseline_regression(df, dep_var, treat_var, control_vars, cov_type='HC3'):
    """
    基准回归
    
    Parameters:
    -----------
    df : DataFrame
        数据
    dep_var : str
        因变量
    treat_var : str
        核心自变量
    control_vars : list
        控制变量列表
    cov_type : str
        标准误类型 ('HC1', 'HC2', 'HC3', 'cluster')
    
    Returns:
    --------
    model : RegressionResults
        回归结果
    n : int
        样本量
    """
    controls_str = ' + '.join(control_vars) if control_vars else ''
    formula = f"{dep_var} ~ {treat_var}"
    if controls_str:
        formula += f" + {controls_str}"
    
    vars_needed = [dep_var, treat_var] + control_vars
    df_reg = df.dropna(subset=vars_needed).copy()
    
    if cov_type == 'cluster':
        raise ValueError("聚类标准误需要使用 cluster_se_regression 函数")
    
    model = ols(formula=formula, data=df_reg).fit(cov_type=cov_type)
    
    return model, len(df_reg)


def fixed_effects_regression(df, dep_var, treat_var, control_vars, fe_vars, cov_type='HC3'):
    """
    固定效应模型
    
    Parameters:
    -----------
    df : DataFrame
        数据
    dep_var : str
        因变量
    treat_var : str
        核心自变量
    control_vars : list
        控制变量列表
    fe_vars : list
        固定效应变量列表
    cov_type : str
        标准误类型
    
    Returns:
    --------
    model : RegressionResults
        回归结果
    n : int
        样本量
    """
    controls_str = ' + '.join(control_vars) if control_vars else ''
    formula = f"{dep_var} ~ {treat_var}"
    if controls_str:
        formula += f" + {controls_str}"
    
    # 添加固定效应
    for fe_var in fe_vars:
        if fe_var in df.columns:
            formula += f" + C({fe_var})"
    
    vars_needed = [dep_var, treat_var] + control_vars + fe_vars
    df_reg = df.dropna(subset=vars_needed).copy()
    
    model = ols(formula=formula, data=df_reg).fit(cov_type=cov_type)
    
    return model, len(df_reg)


def cluster_se_regression(df, dep_var, treat_var, control_vars, cluster_var):
    """
    聚类标准误回归
    
    Parameters:
    -----------
    df : DataFrame
        数据
    dep_var : str
        因变量
    treat_var : str
        核心自变量
    control_vars : list
        控制变量列表
    cluster_var : str
        聚类变量
    
    Returns:
    --------
    model : RegressionResults
        回归结果
    n : int
        样本量
    """
    controls_str = ' + '.join(control_vars) if control_vars else ''
    formula = f"{dep_var} ~ {treat_var}"
    if controls_str:
        formula += f" + {controls_str}"
    
    vars_needed = [dep_var, treat_var] + control_vars + [cluster_var]
    df_reg = df.dropna(subset=vars_needed).copy()
    
    model = ols(formula=formula, data=df_reg).fit(
        cov_type='cluster',
        cov_kwds={'groups': df_reg[cluster_var]}
    )
    
    return model, len(df_reg)


def weighted_least_squares(df, dep_var, treat_var, control_vars, weight_var):
    """
    加权最小二乘法 (WLS)
    
    Parameters:
    -----------
    df : DataFrame
        数据
    dep_var : str
        因变量
    treat_var : str
        核心自变量
    control_vars : list
        控制变量列表
    weight_var : str
        权重变量
    
    Returns:
    --------
    model : RegressionResults
        回归结果
    n : int
        样本量
    """
    controls_str = ' + '.join(control_vars) if control_vars else ''
    formula = f"{dep_var} ~ {treat_var}"
    if controls_str:
        formula += f" + {controls_str}"
    
    vars_needed = [dep_var, treat_var] + control_vars + [weight_var]
    df_reg = df.dropna(subset=vars_needed).copy()
    
    # 确保权重为正
    weights = df_reg[weight_var]
    if weights.min() <= 0:
        weights = weights - weights.min() + 1
    
    model = ols(formula=formula, data=df_reg).fit(weights=weights)
    
    return model, len(df_reg)


# ============================================================
# 样本层面优化
# ============================================================
def subsample_analysis(df, dep_var, treat_var, control_vars, 
                       filter_var, filter_condition):
    """
    子样本分析
    
    Parameters:
    -----------
    df : DataFrame
        数据
    dep_var : str
        因变量
    treat_var : str
        核心自变量
    control_vars : list
        控制变量列表
    filter_var : str
        筛选变量
    filter_condition : callable
        筛选条件函数
    
    Returns:
    --------
    model : RegressionResults
        回归结果
    n : int
        样本量
    """
    sub_df = df[filter_condition(df[filter_var])].copy()
    
    if len(sub_df) < 50:
        return None, len(sub_df)
    
    model, n = baseline_regression(sub_df, dep_var, treat_var, control_vars)
    
    return model, n


def exclude_outliers(df, dep_var, treat_var, control_vars, 
                     method='iqr', threshold=1.5):
    """
    剔除异常值
    
    Parameters:
    -----------
    df : DataFrame
        数据
    dep_var : str
        因变量
    treat_var : str
        核心自变量
    control_vars : list
        控制变量列表
    method : str
        方法 ('iqr', 'zscore', 'percentile')
    threshold : float
        阈值
    
    Returns:
    --------
    model : RegressionResults
        回归结果
    n : int
        样本量
    """
    df_clean = df.copy()
    
    if method == 'iqr':
        Q1 = df_clean[dep_var].quantile(0.25)
        Q3 = df_clean[dep_var].quantile(0.75)
        IQR = Q3 - Q1
        lower = Q1 - threshold * IQR
        upper = Q3 + threshold * IQR
        df_clean = df_clean[(df_clean[dep_var] >= lower) & (df_clean[dep_var] <= upper)]
    
    elif method == 'zscore':
        z_scores = np.abs((df_clean[dep_var] - df_clean[dep_var].mean()) / df_clean[dep_var].std())
        df_clean = df_clean[z_scores < threshold]
    
    elif method == 'percentile':
        lower = df_clean[dep_var].quantile(threshold / 100)
        upper = df_clean[dep_var].quantile(1 - threshold / 100)
        df_clean = df_clean[(df_clean[dep_var] >= lower) & (df_clean[dep_var] <= upper)]
    
    if len(df_clean) < 50:
        return None, len(df_clean)
    
    model, n = baseline_regression(df_clean, dep_var, treat_var, control_vars)
    
    return model, n


# ============================================================
# 诊断检验
# ============================================================
def calculate_vif(df, variables):
    """
    计算方差膨胀因子 (VIF)
    
    Parameters:
    -----------
    df : DataFrame
        数据
    variables : list
        变量列表
    
    Returns:
    --------
    vif_data : DataFrame
        VIF 结果
    """
    from statsmodels.stats.outliers_influence import variance_inflation_factor
    
    df_vif = df[variables].dropna()
    
    if len(df_vif) < len(variables) + 1:
        return pd.DataFrame({
            'Variable': variables,
            'VIF': [np.nan] * len(variables),
            'VIF_Category': ['Insufficient data'] * len(variables)
        })
    
    vif_data = pd.DataFrame()
    vif_data["Variable"] = variables
    vif_data["VIF"] = [variance_inflation_factor(df_vif.values, i) 
                       for i in range(len(variables))]
    vif_data["VIF_Category"] = vif_data["VIF"].apply(
        lambda x: "Low (<5)" if x < 5 else ("Moderate (5-10)" if x < 10 else "High (>10)")
    )
    
    return vif_data


def breusch_pagan_test(model):
    """
    BP 异方差检验
    
    Parameters:
    -----------
    model : RegressionResults
        回归模型
    
    Returns:
    --------
    dict : 检验结果
    """
    import statsmodels.api as sm
    
    bp_test = sm.stats.diagnostic.het_breuschpagan(model.resid, model.model.exog)
    
    return {
        'lm_statistic': bp_test[0],
        'lm_pvalue': bp_test[1],
        'f_statistic': bp_test[2],
        'f_pvalue': bp_test[3],
        'conclusion': '存在异方差 (p<0.05)' if bp_test[1] < 0.05 else '不存在异方差 (p>=0.05)'
    }


def durbin_watson_test(model):
    """
    DW 自相关检验
    
    Parameters:
    -----------
    model : RegressionResults
        回归模型
    
    Returns:
    --------
    dict : 检验结果
    """
    import statsmodels.api as sm
    
    dw_stat = sm.stats.durbin_watson(model.resid)
    
    if dw_stat < 1.5:
        conclusion = '可能存在正自相关'
    elif dw_stat > 2.5:
        conclusion = '可能存在负自相关'
    else:
        conclusion = '无明显自相关'
    
    return {
        'dw_statistic': dw_stat,
        'conclusion': conclusion
    }


def ramsey_reset_test(model, n_fitted=2):
    """
    RESET 检验（模型设定错误）
    
    Parameters:
    -----------
    model : RegressionResults
        回归模型
    n_fitted : int
        添加的拟合值幂次项数
    
    Returns:
    --------
    dict : 检验结果
    """
    import statsmodels.api as sm
    
    # 获取拟合值
    fitted = model.fittedvalues
    
    # 构建新公式
    new_regressors = np.column_stack([fitted ** i for i in range(2, n_fitted + 2)])
    
    # RESET 检验
    reset_test = sm.stats.diagnostic.linear_reset(model, power=n_fitted)
    
    # 处理不同版本的 statsmodels 返回值
    if hasattr(reset_test, 'statistic') and hasattr(reset_test, 'pvalue'):
        # 新版本返回 ContrastResults 对象
        f_stat = reset_test.statistic
        p_val = reset_test.pvalue
    else:
        # 旧版本返回元组
        f_stat = reset_test[0]
        p_val = reset_test[1]
    
    return {
        'f_statistic': f_stat,
        'pvalue': p_val,
        'conclusion': '模型设定可能有问题 (p<0.05)' if p_val < 0.05 else '模型设定合理 (p>=0.05)'
    }


# ============================================================
# 完整优化流程
# ============================================================
def significance_optimization(df, dep_var, treat_var, control_vars,
                               target_p=0.05, output_dir=None):
    """
    完整显著性优化流程
    
    Parameters:
    -----------
    df : DataFrame
        数据
    dep_var : str
        因变量
    treat_var : str
        核心自变量
    control_vars : list
        控制变量列表
    target_p : float
        目标显著性水平
    output_dir : str
        输出目录
    
    Returns:
    --------
    results : dict
        所有优化结果
    """
    print("=" * 60)
    print("开始显著性优化分析...")
    print("=" * 60)
    
    results = {}
    optimization_log = []
    
    # ========================================================
    # 阶段 1: 基准回归
    # ========================================================
    print("\n" + "=" * 60)
    print("【阶段 1】基准回归")
    print("=" * 60)
    
    base_model, base_n = baseline_regression(df, dep_var, treat_var, control_vars)
    base_coef = base_model.params[treat_var]
    base_p = base_model.pvalues[treat_var]
    
    results['baseline'] = {
        'model': base_model,
        'n': base_n,
        'coef': base_coef,
        'pvalue': base_p,
        'significant': base_p < target_p
    }
    
    print(f"基准回归：系数={base_coef:.4f}, p={base_p:.4f} {'***' if base_p < 0.01 else ('**' if base_p < 0.05 else ('*' if base_p < 0.1 else ''))}")
    optimization_log.append({
        'method': '基准回归',
        'coef': base_coef,
        'se': base_model.bse[treat_var],
        'pvalue': base_p,
        'n': base_n
    })
    
    # ========================================================
    # 阶段 2: 数据层面优化
    # ========================================================
    print("\n" + "=" * 60)
    print("【阶段 2】数据层面优化")
    print("=" * 60)
    
    # 2.1 Winsorize 缩尾 (1%)
    continuous_vars = [dep_var] + control_vars
    df_winsorized_1 = winsorize_df(df, continuous_vars, limits=(0.01, 0.99))
    model_w1, n_w1 = baseline_regression(df_winsorized_1, dep_var, treat_var, control_vars)
    p_w1 = model_w1.pvalues[treat_var]
    
    results['winsorized_1pct'] = {
        'model': model_w1,
        'n': n_w1,
        'coef': model_w1.params[treat_var],
        'pvalue': p_w1,
        'significant': p_w1 < target_p
    }
    
    print(f"Winsorize 1% 缩尾：系数={model_w1.params[treat_var]:.4f}, p={p_w1:.4f}")
    optimization_log.append({
        'method': 'Winsorize 1% 缩尾',
        'coef': model_w1.params[treat_var],
        'se': model_w1.bse[treat_var],
        'pvalue': p_w1,
        'n': n_w1
    })
    
    # 2.2 Winsorize 缩尾 (5%)
    df_winsorized_5 = winsorize_df(df, continuous_vars, limits=(0.05, 0.95))
    model_w5, n_w5 = baseline_regression(df_winsorized_5, dep_var, treat_var, control_vars)
    p_w5 = model_w5.pvalues[treat_var]
    
    results['winsorized_5pct'] = {
        'model': model_w5,
        'n': n_w5,
        'coef': model_w5.params[treat_var],
        'pvalue': p_w5,
        'significant': p_w5 < target_p
    }
    
    print(f"Winsorize 5% 缩尾：系数={model_w5.params[treat_var]:.4f}, p={p_w5:.4f}")
    optimization_log.append({
        'method': 'Winsorize 5% 缩尾',
        'coef': model_w5.params[treat_var],
        'se': model_w5.bse[treat_var],
        'pvalue': p_w5,
        'n': n_w5
    })
    
    # 2.3 对数转换
    if df[dep_var].min() > 0:
        df_log = df.copy()
        df_log[f'{dep_var}_log'] = np.log(df_log[dep_var])
        model_log, n_log = baseline_regression(df_log, f'{dep_var}_log', treat_var, control_vars)
        p_log = model_log.pvalues[treat_var]
        
        results['log_transformed'] = {
            'model': model_log,
            'n': n_log,
            'coef': model_log.params[treat_var],
            'pvalue': p_log,
            'significant': p_log < target_p
        }
        
        print(f"对数转换：系数={model_log.params[treat_var]:.4f}, p={p_log:.4f}")
        optimization_log.append({
            'method': '对数转换',
            'coef': model_log.params[treat_var],
            'se': model_log.bse[treat_var],
            'pvalue': p_log,
            'n': n_log
        })
    
    # ========================================================
    # 阶段 3: 模型层面优化
    # ========================================================
    print("\n" + "=" * 60)
    print("【阶段 3】模型层面优化")
    print("=" * 60)
    
    # 3.1 不同稳健标准误
    for hc in ['HC1', 'HC2', 'HC3']:
        model_hc, n_hc = baseline_regression(df, dep_var, treat_var, control_vars, cov_type=hc)
        p_hc = model_hc.pvalues[treat_var]
        
        results[f'robust_{hc}'] = {
            'model': model_hc,
            'n': n_hc,
            'coef': model_hc.params[treat_var],
            'pvalue': p_hc,
            'significant': p_hc < target_p
        }
        
        print(f"稳健标准误 ({hc}): 系数={model_hc.params[treat_var]:.4f}, p={p_hc:.4f}")
        optimization_log.append({
            'method': f'稳健标准误 ({hc})',
            'coef': model_hc.params[treat_var],
            'se': model_hc.bse[treat_var],
            'pvalue': p_hc,
            'n': n_hc
        })
    
    # 3.2 固定效应
    if 'region_old' in df.columns:
        model_region_fe, n_fe = fixed_effects_regression(df, dep_var, treat_var, control_vars, ['region_old'])
        p_fe = model_region_fe.pvalues[treat_var]
        
        results['region_fe'] = {
            'model': model_region_fe,
            'n': n_fe,
            'coef': model_region_fe.params[treat_var],
            'pvalue': p_fe,
            'significant': p_fe < target_p
        }
        
        print(f"地区固定效应：系数={model_region_fe.params[treat_var]:.4f}, p={p_fe:.4f}")
        optimization_log.append({
            'method': '地区固定效应',
            'coef': model_region_fe.params[treat_var],
            'se': model_region_fe.bse[treat_var],
            'pvalue': p_fe,
            'n': n_fe
        })
    
    if 'year' in df.columns:
        model_year_fe, n_fe = fixed_effects_regression(df, dep_var, treat_var, control_vars, ['year'])
        p_fe = model_year_fe.pvalues[treat_var]
        
        results['year_fe'] = {
            'model': model_year_fe,
            'n': n_fe,
            'coef': model_year_fe.params[treat_var],
            'pvalue': p_fe,
            'significant': p_fe < target_p
        }
        
        print(f"年份固定效应：系数={model_year_fe.params[treat_var]:.4f}, p={p_fe:.4f}")
        optimization_log.append({
            'method': '年份固定效应',
            'coef': model_year_fe.params[treat_var],
            'se': model_year_fe.bse[treat_var],
            'pvalue': p_fe,
            'n': n_fe
        })
        
        if 'region_old' in df.columns:
            model_twfe, n_twfe = fixed_effects_regression(df, dep_var, treat_var, control_vars, ['region_old', 'year'])
            p_twfe = model_twfe.pvalues[treat_var]
            
            results['twfe'] = {
                'model': model_twfe,
                'n': n_twfe,
                'coef': model_twfe.params[treat_var],
                'pvalue': p_twfe,
                'significant': p_twfe < target_p
            }
            
            print(f"双向固定效应：系数={model_twfe.params[treat_var]:.4f}, p={p_twfe:.4f}")
            optimization_log.append({
                'method': '双向固定效应',
                'coef': model_twfe.params[treat_var],
                'se': model_twfe.bse[treat_var],
                'pvalue': p_twfe,
                'n': n_twfe
            })
    
    # 3.3 聚类标准误
    for cluster_var in ['city', 'region_old', 'householdID']:
        if cluster_var in df.columns:
            try:
                model_cluster, n_cluster = cluster_se_regression(df, dep_var, treat_var, control_vars, cluster_var)
                p_cluster = model_cluster.pvalues[treat_var]
                
                results[f'cluster_{cluster_var}'] = {
                    'model': model_cluster,
                    'n': n_cluster,
                    'coef': model_cluster.params[treat_var],
                    'pvalue': p_cluster,
                    'significant': p_cluster < target_p
                }
                
                print(f"聚类标准误 ({cluster_var}): 系数={model_cluster.params[treat_var]:.4f}, p={p_cluster:.4f}")
                optimization_log.append({
                    'method': f'聚类标准误 ({cluster_var})',
                    'coef': model_cluster.params[treat_var],
                    'se': model_cluster.bse[treat_var],
                    'pvalue': p_cluster,
                    'n': n_cluster
                })
            except:
                pass
    
    # ========================================================
    # 阶段 4: 样本层面优化
    # ========================================================
    print("\n" + "=" * 60)
    print("【阶段 4】样本层面优化")
    print("=" * 60)
    
    # 4.1 剔除因变量极端值
    model_exclude, n_exclude = exclude_outliers(df, dep_var, treat_var, control_vars, 
                                                 method='percentile', threshold=1)
    if model_exclude is not None:
        p_exclude = model_exclude.pvalues[treat_var]
        
        results['exclude_outliers_1pct'] = {
            'model': model_exclude,
            'n': n_exclude,
            'coef': model_exclude.params[treat_var],
            'pvalue': p_exclude,
            'significant': p_exclude < target_p
        }
        
        print(f"剔除 1% 极端值：系数={model_exclude.params[treat_var]:.4f}, p={p_exclude:.4f}, N={n_exclude}")
        optimization_log.append({
            'method': '剔除 1% 极端值',
            'coef': model_exclude.params[treat_var],
            'se': model_exclude.bse[treat_var],
            'pvalue': p_exclude,
            'n': n_exclude
        })
    
    # ========================================================
    # 阶段 5: 诊断检验
    # ========================================================
    print("\n" + "=" * 60)
    print("【阶段 5】诊断检验")
    print("=" * 60)
    
    diagnostics = {}
    
    # VIF
    vif_vars = [treat_var] + control_vars[:5]
    vif_result = calculate_vif(df, vif_vars)
    diagnostics['vif'] = vif_result
    print("\nVIF 检验:")
    print(vif_result.to_string())
    
    # BP 检验
    bp_result = breusch_pagan_test(base_model)
    diagnostics['bp_test'] = bp_result
    print(f"\nBP 异方差检验：{bp_result['conclusion']}")
    
    # DW 检验
    dw_result = durbin_watson_test(base_model)
    diagnostics['dw_test'] = dw_result
    print(f"DW 自相关检验：{dw_result['conclusion']} (DW={dw_result['dw_statistic']:.4f})")
    
    # RESET 检验
    reset_result = ramsey_reset_test(base_model)
    diagnostics['reset_test'] = reset_result
    print(f"RESET 检验：{reset_result['conclusion']}")
    
    # ========================================================
    # 生成报告
    # ========================================================
    if output_dir is None:
        output_dir = os.getcwd()
    
    results_dict = {
        'optimization_log': optimization_log,
        'diagnostics': diagnostics,
        'best_model': None
    }
    
    # 找到最显著的结果
    significant_results = [(k, v) for k, v in results.items() if v['significant']]
    if significant_results:
        best_key, best_val = min(significant_results, key=lambda x: x[1]['pvalue'])
        results_dict['best_model'] = {
            'method': best_key,
            'coef': best_val['coef'],
            'pvalue': best_val['pvalue']
        }
    
    # 生成 Word 报告
    output_path = os.path.join(output_dir, '显著性优化报告.docx')
    generate_optimization_report(results_dict, results, output_path)
    print(f"\n[OK] Word 报告已保存：{output_path}")
    
    print("\n" + "=" * 60)
    print("显著性优化分析完成！")
    print("=" * 60)
    
    # 总结
    print("\n【优化总结】")
    baseline_sig = "***" if base_p < 0.01 else ("**" if base_p < 0.05 else ("*" if base_p < 0.1 else "不显著"))
    print(f"基准回归：p={base_p:.4f} ({baseline_sig})")
    
    if results_dict['best_model']:
        best = results_dict['best_model']
        best_sig = "***" if best['pvalue'] < 0.01 else ("**" if best['pvalue'] < 0.05 else ("*" if best['pvalue'] < 0.1 else ""))
        print(f"最优方法：{best['method']}")
        print(f"最优结果：p={best['pvalue']:.4f} ({best_sig})")
    
    return results_dict, results


# ============================================================
# Word 报告生成
# ============================================================
def generate_optimization_report(results_dict, all_results, output_path):
    """
    生成 Word 格式优化报告
    
    Parameters:
    -----------
    results_dict : dict
        结果字典
    all_results : dict
        所有优化结果
    output_path : str
        输出文件路径
    """
    doc = Document()
    
    # 标题
    title = doc.add_heading('显著性优化分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加日期
    date_para = doc.add_paragraph(f'分析日期：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 1. 优化结果对比
    doc.add_heading('一、优化结果对比', level=1)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    headers = ['方法', '系数', '标准误', 'P 值', '显著性', '样本量']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    optimization_log = results_dict.get('optimization_log', [])
    for item in optimization_log:
        row = table.add_row().cells
        row[0].text = item['method']
        row[1].text = f"{item['coef']:.4f}"
        row[2].text = f"({item['se']:.4f})"
        p = item['pvalue']
        sig = '***' if p < 0.01 else ('**' if p < 0.05 else ('*' if p < 0.1 else ''))
        row[3].text = f"{p:.4f}"
        row[4].text = sig
        row[5].text = str(item['n'])
    
    doc.add_paragraph('\n注：*** p<0.01, ** p<0.05, * p<0.1')
    
    # 2. 最佳模型
    if results_dict.get('best_model'):
        doc.add_heading('二、最优模型', level=1)
        best = results_dict['best_model']
        doc.add_paragraph(f'方法：{best["method"]}')
        doc.add_paragraph(f'系数：{best["coef"]:.4f}')
        doc.add_paragraph(f'P 值：{best["pvalue"]:.4f}')
    
    # 3. 诊断检验
    if 'diagnostics' in results_dict:
        doc.add_heading('三、诊断检验', level=1)
        diag = results_dict['diagnostics']
        
        if 'vif' in diag:
            doc.add_heading('3.1 多重共线性检验 (VIF)', level=2)
            vif_table = doc.add_table(rows=1, cols=3)
            vif_table.style = 'Table Grid'
            vif_headers = ['变量', 'VIF 值', '程度']
            for i, header in enumerate(vif_headers):
                vif_table.rows[0].cells[i].text = header
                vif_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            
            for _, row_data in diag['vif'].iterrows():
                vif_row = vif_table.add_row().cells
                vif_row[0].text = str(row_data['Variable'])
                vif_row[1].text = f"{row_data['VIF']:.2f}" if not np.isnan(row_data['VIF']) else 'N/A'
                vif_row[2].text = str(row_data['VIF_Category'])
        
        if 'bp_test' in diag:
            doc.add_heading('3.2 异方差检验 (Breusch-Pagan)', level=2)
            bp = diag['bp_test']
            doc.add_paragraph(f'LM 统计量：{bp["lm_statistic"]:.4f}')
            doc.add_paragraph(f'P 值：{bp["lm_pvalue"]:.4f}')
            doc.add_paragraph(f'结论：{bp["conclusion"]}')
        
        if 'dw_test' in diag:
            doc.add_heading('3.3 自相关检验 (Durbin-Watson)', level=2)
            dw = diag['dw_test']
            doc.add_paragraph(f'DW 统计量：{dw["dw_statistic"]:.4f}')
            doc.add_paragraph(f'结论：{dw["conclusion"]}')
        
        if 'reset_test' in diag:
            doc.add_heading('3.4 模型设定检验 (RESET)', level=2)
            reset = diag['reset_test']
            doc.add_paragraph(f'F 统计量：{reset["f_statistic"]:.4f}')
            doc.add_paragraph(f'P 值：{reset["pvalue"]:.4f}')
            doc.add_paragraph(f'结论：{reset["conclusion"]}')
    
    # 4. 学术诚信声明
    doc.add_heading('四、学术诚信声明', level=1)
    doc.add_paragraph('''
本分析报告仅供学术研究和教学参考使用。

正当使用：
- 探索不同模型设定下的结果稳健性
- 理解哪些因素会影响统计显著性
- 进行敏感性分析和稳健性检验

不当使用：
- P-hacking（p 值操纵）
- 数据窥探（data dredging）
- 选择性报告结果
- 篡改原始数据

请在学术诚信的前提下使用本工具。
''')
    
    # 保存
    doc.save(output_path)
    return output_path


# ============================================================
# 主函数（命令行调用）
# ============================================================
if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("用法：python optimization_analysis.py <数据文件路径> <因变量> <自变量> [控制变量 1,控制变量 2,...]")
        sys.exit(1)
    
    file_path = sys.argv[1]
    dep_var = sys.argv[2]
    treat_var = sys.argv[3]
    control_vars = sys.argv[4].split(',') if len(sys.argv) > 4 else []
    
    # 读取数据
    if file_path.endswith('.dta'):
        df = pd.read_stata(file_path)
    elif file_path.endswith('.xlsx'):
        df = pd.read_excel(file_path)
    elif file_path.endswith('.csv'):
        df = pd.read_csv(file_path)
    else:
        print(f"不支持的文件格式")
        sys.exit(1)
    
    results, all_results = significance_optimization(df, dep_var, treat_var, control_vars)
