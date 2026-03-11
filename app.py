# -*- coding: utf-8 -*-
"""
一键实证分析 - Web 可视化界面
基于 Streamlit 构建，支持数据上传、变量选择、多元分析、报告下载
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import statsmodels.api as sm
from statsmodels.formula.api import ols
from statsmodels.stats.outliers_influence import variance_inflation_factor
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
import os
import tempfile
import warnings
warnings.filterwarnings('ignore')

# 导入显著性优化模块
try:
    from optimization_analysis import significance_optimization, generate_optimization_report
    OPTIMIZATION_AVAILABLE = True
except ImportError:
    OPTIMIZATION_AVAILABLE = False

# ============================================================
# 页面配置
# ============================================================
st.set_page_config(
    page_title="一键实证分析",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================
# 自定义样式
# ============================================================
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        color: #1E88E5;
        text-align: center;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.5rem;
        font-weight: 600;
        color: #424242;
        margin-top: 1.5rem;
        margin-bottom: 1rem;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 0.5rem;
        color: white;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        padding-left: 20px;
        padding-right: 20px;
        font-weight: 600;
    }
    .success-message {
        padding: 1rem;
        background-color: #d4edda;
        border-color: #c3e6cb;
        border-radius: 0.5rem;
        color: #155724;
    }
    .info-box {
        padding: 1rem;
        background-color: #e3f2fd;
        border-left: 4px solid #2196F3;
        border-radius: 0.25rem;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================
# 变量自动识别模式
# ============================================================
VARIABLE_PATTERNS = {
    'dep_var': ['y', 'outcome', 'depvar', 'result', 'depression', 'iadl', 'health', 
                'income', 'consumption', 'aczb', 'aczhenb', 'acxb', 'numcds', 'social'],
    'treat_var': ['treat', 'treatment', 'D', 'T', 'policy', 'program', 'treat2', 'treat3', 
                  'fv', 'lv', 'insur', 'yanglao_insur'],
    'control_var': ['age', 'gender', 'edu', 'education', 'income', 'married', 'urban', 
                    'familysize', 'lntotalincome', 'lnpergdp', 'numchild', 'rate_laborage'],
    'id_var': ['id', 'householdid'],
    'time_var': ['year', 'time', 'period', 'wave'],
    'region_var': ['region', 'province', 'city', 'county', 'region_old', 'lev_inc']
}

# ============================================================
# 核心功能函数
# ============================================================

def auto_identify_variables(df):
    """自动识别变量类型"""
    identified = {
        'dep_var': None,
        'treat_var': None,
        'control_vars': [],
        'id_var': None,
        'time_var': None,
        'region_var': None
    }
    
    for var_type, patterns in VARIABLE_PATTERNS.items():
        if var_type == 'control_var':
            for pattern in patterns:
                for col in df.columns:
                    if pattern.lower() in col.lower() and col not in identified['control_vars']:
                        if col not in [identified['dep_var'], identified['treat_var']]:
                            identified['control_vars'].append(col)
        else:
            for pattern in patterns:
                for col in df.columns:
                    if pattern.lower() == col.lower() or pattern.lower() in col.lower():
                        if identified[var_type] is None:
                            identified[var_type] = col
                        break
    
    return identified


def read_data(file_path, file_type):
    """读取数据文件"""
    if file_type == '.dta':
        df = pd.read_stata(file_path)
    elif file_type in ['.xlsx', '.xls']:
        df = pd.read_excel(file_path)
    elif file_type == '.csv':
        df = pd.read_csv(file_path, encoding='utf-8')
    elif file_type == '.sav':
        df = pd.read_spss(file_path)
    else:
        raise ValueError(f"不支持的文件格式：{file_type}")
    return df


def descriptive_statistics(df, variables=None):
    """生成描述性统计"""
    if variables is None:
        variables = df.select_dtypes(include=[np.number]).columns.tolist()
    
    desc = df[variables].describe()
    desc.loc['skewness'] = df[variables].skew()
    desc.loc['kurtosis'] = df[variables].kurtosis()
    desc.loc['missing'] = df[variables].isnull().sum()
    desc.loc['missing_pct'] = (df[variables].isnull().sum() / len(df) * 100).round(2)
    
    return desc


def baseline_regression(df, dep_var, treat_var, control_vars, robust_se=True):
    """基准回归分析"""
    controls_str = ' + '.join(control_vars) if control_vars else ''
    formula = f"{dep_var} ~ {treat_var}"
    if controls_str:
        formula += f" + {controls_str}"
    
    vars_needed = [dep_var, treat_var] + control_vars
    df_reg = df.dropna(subset=vars_needed).copy()
    
    cov_type = 'HC3' if robust_se else 'nonrobust'
    model = ols(formula=formula, data=df_reg).fit(cov_type=cov_type)
    
    return model, df_reg


def fixed_effects_regression(df, dep_var, treat_var, control_vars, fe_vars):
    """固定效应模型"""
    controls_str = ' + '.join(control_vars) if control_vars else ''
    formula = f"{dep_var} ~ {treat_var}"
    if controls_str:
        formula += f" + {controls_str}"
    
    for fe_var in fe_vars:
        if fe_var in df.columns:
            formula += f" + C({fe_var})"
    
    vars_needed = [dep_var, treat_var] + control_vars + fe_vars
    df_reg = df.dropna(subset=vars_needed).copy()
    
    model = ols(formula=formula, data=df_reg).fit(cov_type='HC3')
    
    return model, df_reg


def heterogeneity_analysis(df, dep_var, treat_var, control_vars, group_var):
    """异质性分析"""
    results = []
    unique_groups = df[group_var].dropna().unique()
    
    for group_val in unique_groups:
        sub_df = df[df[group_var] == group_val]
        
        if len(sub_df) < 50:
            continue
        
        try:
            model, _ = baseline_regression(sub_df, dep_var, treat_var, control_vars)
            results.append({
                'group': str(group_val),
                'n': len(sub_df),
                'coef': model.params[treat_var],
                'se': model.bse[treat_var],
                'pvalue': model.pvalues[treat_var],
                'r2': model.rsquared
            })
        except:
            continue
    
    return pd.DataFrame(results)


def winsorize_series(series, limits=(0.01, 0.99)):
    """Winsorize 缩尾处理"""
    lower = series.quantile(limits[0])
    upper = series.quantile(limits[1])
    return series.clip(lower=lower, upper=upper)


def robustness_checks(df, dep_var, treat_var, control_vars):
    """稳健性检验"""
    results = []
    
    # 1. Winsorize 1%
    continuous_vars = [dep_var] + control_vars
    df_winsorized = df.copy()
    for var in continuous_vars:
        if var in df_winsorized.columns and df_winsorized[var].dtype in [np.float64, np.int64]:
            df_winsorized[var] = winsorize_series(df_winsorized[var], limits=(0.01, 0.99))
    
    model_w1, _ = baseline_regression(df_winsorized, dep_var, treat_var, control_vars)
    results.append({
        'method': 'Winsorize 1%',
        'coef': model_w1.params[treat_var],
        'se': model_w1.bse[treat_var],
        'pvalue': model_w1.pvalues[treat_var]
    })
    
    # 2. Winsorize 5%
    df_winsorized_5 = df.copy()
    for var in continuous_vars:
        if var in df_winsorized_5.columns and df_winsorized_5[var].dtype in [np.float64, np.int64]:
            df_winsorized_5[var] = winsorize_series(df_winsorized_5[var], limits=(0.05, 0.95))
    
    model_w5, _ = baseline_regression(df_winsorized_5, dep_var, treat_var, control_vars)
    results.append({
        'method': 'Winsorize 5%',
        'coef': model_w5.params[treat_var],
        'se': model_w5.bse[treat_var],
        'pvalue': model_w5.pvalues[treat_var]
    })
    
    # 3. 不同稳健标准误
    for hc in ['HC1', 'HC2', 'HC3']:
        controls_str = ' + '.join(control_vars) if control_vars else ''
        formula = f"{dep_var} ~ {treat_var}"
        if controls_str:
            formula += f" + {controls_str}"
        vars_needed = [dep_var, treat_var] + control_vars
        df_reg = df.dropna(subset=vars_needed).copy()
        model = ols(formula=formula, data=df_reg).fit(cov_type=hc)
        results.append({
            'method': f'稳健标准误 {hc}',
            'coef': model.params[treat_var],
            'se': model.bse[treat_var],
            'pvalue': model.pvalues[treat_var]
        })
    
    return pd.DataFrame(results)


def calculate_vif(df, variables):
    """计算方差膨胀因子"""
    df_vif = df[variables].dropna()
    
    vif_data = pd.DataFrame()
    vif_data["变量"] = variables
    vif_data["VIF"] = [variance_inflation_factor(df_vif.values, i) 
                       for i in range(len(variables))]
    vif_data["程度"] = vif_data["VIF"].apply(
        lambda x: "低" if x < 5 else ("中等" if x < 10 else "高")
    )
    
    return vif_data


def breusch_pagan_test(model):
    """BP 异方差检验"""
    bp_test = sm.stats.diagnostic.het_breuschpagan(model.resid, model.model.exog)
    return {
        'lm_statistic': bp_test[0],
        'lm_pvalue': bp_test[1],
        'f_statistic': bp_test[2],
        'f_pvalue': bp_test[3],
        'conclusion': '存在异方差' if bp_test[1] < 0.05 else '不存在异方差'
    }


def durbin_watson_test(model):
    """DW 自相关检验"""
    dw_stat = sm.stats.durbin_watson(model.resid)
    
    if dw_stat < 1.5:
        conclusion = '可能存在正自相关'
    elif dw_stat > 2.5:
        conclusion = '可能存在负自相关'
    else:
        conclusion = '无明显自相关'
    
    return {
        'dw_statistic': dw_stat,
        'conclusion': conclusion
    }


def generate_word_report(results_dict, title='实证分析报告'):
    """生成 Word 格式分析报告"""
    doc = Document()
    
    # 标题
    doc_heading = doc.add_heading(title, 0)
    doc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加日期
    from datetime import datetime
    date_para = doc.add_paragraph(f'分析日期：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 1. 数据描述
    if 'data_info' in results_dict:
        doc.add_heading('一、数据描述', level=1)
        data_info = results_dict['data_info']
        doc.add_paragraph(f'样本量：{data_info["n_obs"]:,} 个观测值')
        doc.add_paragraph(f'变量数：{data_info["n_vars"]} 个变量')
    
    # 2. 描述性统计
    if 'descriptive' in results_dict:
        doc.add_heading('二、描述性统计', level=1)
        desc = results_dict['descriptive']
        
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        headers = ['变量', '均值', '标准差', '最小值', '25%', '中位数', '最大值']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for var in desc.columns[:15]:  # 限制前15个变量
            row = table.add_row().cells
            row[0].text = str(var)
            row[1].text = f"{desc[var]['mean']:.4f}" if 'mean' in desc.index else 'N/A'
            row[2].text = f"{desc[var]['std']:.4f}" if 'std' in desc.index else 'N/A'
            row[3].text = f"{desc[var]['min']:.4f}" if 'min' in desc.index else 'N/A'
            row[4].text = f"{desc[var]['25%']:.4f}" if '25%' in desc.index else 'N/A'
            row[5].text = f"{desc[var]['50%']:.4f}" if '50%' in desc.index else 'N/A'
            row[6].text = f"{desc[var]['max']:.4f}" if 'max' in desc.index else 'N/A'
    
    # 3. 基准回归
    if 'baseline' in results_dict:
        doc.add_heading('三、基准回归结果', level=1)
        model = results_dict['baseline']['model']
        n = results_dict['baseline']['n']
        
        doc.add_paragraph(f'样本量：N = {n:,}')
        
        reg_table = doc.add_table(rows=1, cols=4)
        reg_table.style = 'Table Grid'
        reg_headers = ['变量', '系数', '标准误', 'P 值']
        for i, header in enumerate(reg_headers):
            reg_table.rows[0].cells[i].text = header
            reg_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for var_name in model.params.index:
            coef = model.params[var_name]
            se = model.bse[var_name]
            p = model.pvalues[var_name]
            sig = '***' if p < 0.01 else ('**' if p < 0.05 else ('*' if p < 0.1 else ''))
            
            row = reg_table.add_row().cells
            row[0].text = str(var_name)
            row[1].text = f"{coef:.4f}{sig}"
            row[2].text = f"({se:.4f})"
            row[3].text = f"{p:.4f}"
        
        doc.add_paragraph('\n注：*** p<0.01, ** p<0.05, * p<0.1')
        doc.add_paragraph(f'R² = {model.rsquared:.4f}, 调整后 R² = {model.rsquared_adj:.4f}')
    
    # 4. 异质性分析
    if 'heterogeneity' in results_dict and results_dict['heterogeneity'] is not None:
        doc.add_heading('四、异质性分析', level=1)
        het_df = results_dict['heterogeneity']
        
        het_table = doc.add_table(rows=1, cols=5)
        het_table.style = 'Table Grid'
        het_headers = ['组别', '样本量', '系数', '标准误', 'P 值']
        for i, header in enumerate(het_headers):
            het_table.rows[0].cells[i].text = header
            het_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for _, row_data in het_df.iterrows():
            row = het_table.add_row().cells
            row[0].text = str(row_data['group'])
            row[1].text = str(int(row_data['n']))
            row[2].text = f"{row_data['coef']:.4f}"
            row[3].text = f"({row_data['se']:.4f})"
            row[4].text = f"{row_data['pvalue']:.4f}"
    
    # 5. 稳健性检验
    if 'robustness' in results_dict and results_dict['robustness'] is not None:
        doc.add_heading('五、稳健性检验', level=1)
        robust_df = results_dict['robustness']
        
        robust_table = doc.add_table(rows=1, cols=4)
        robust_table.style = 'Table Grid'
        robust_headers = ['检验方法', '系数', '标准误', 'P 值']
        for i, header in enumerate(robust_headers):
            robust_table.rows[0].cells[i].text = header
            robust_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for _, row_data in robust_df.iterrows():
            sig = '***' if row_data['pvalue'] < 0.01 else ('**' if row_data['pvalue'] < 0.05 else ('*' if row_data['pvalue'] < 0.1 else ''))
            row = robust_table.add_row().cells
            row[0].text = str(row_data['method'])
            row[1].text = f"{row_data['coef']:.4f}{sig}"
            row[2].text = f"({row_data['se']:.4f})"
            row[3].text = f"{row_data['pvalue']:.4f}"
    
    # 6. 诊断检验
    if 'diagnostics' in results_dict:
        doc.add_heading('六、诊断检验', level=1)
        diag = results_dict['diagnostics']
        
        if 'vif' in diag:
            doc.add_heading('6.1 多重共线性检验 (VIF)', level=2)
            vif_table = doc.add_table(rows=1, cols=3)
            vif_table.style = 'Table Grid'
            vif_headers = ['变量', 'VIF 值', '程度']
            for i, header in enumerate(vif_headers):
                vif_table.rows[0].cells[i].text = header
                vif_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            
            for _, row_data in diag['vif'].iterrows():
                row = vif_table.add_row().cells
                row[0].text = str(row_data['变量'])
                row[1].text = f"{row_data['VIF']:.2f}"
                row[2].text = row_data['程度']
        
        if 'bp_test' in diag:
            doc.add_heading('6.2 异方差检验 (Breusch-Pagan)', level=2)
            bp = diag['bp_test']
            doc.add_paragraph(f'LM 统计量：{bp["lm_statistic"]:.4f}')
            doc.add_paragraph(f'P 值：{bp["lm_pvalue"]:.4f}')
            doc.add_paragraph(f'结论：{bp["conclusion"]}')
        
        if 'dw_test' in diag:
            doc.add_heading('6.3 自相关检验 (Durbin-Watson)', level=2)
            dw = diag['dw_test']
            doc.add_paragraph(f'DW 统计量：{dw["dw_statistic"]:.4f}')
            doc.add_paragraph(f'结论：{dw["conclusion"]}')
    
    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def format_regression_table(model, treat_var=None):
    """格式化回归结果表格"""
    data = []
    for var_name in model.params.index:
        coef = model.params[var_name]
        se = model.bse[var_name]
        t = model.tvalues[var_name]
        p = model.pvalues[var_name]
        
        if p < 0.01:
            sig = '***'
        elif p < 0.05:
            sig = '**'
        elif p < 0.1:
            sig = '*'
        else:
            sig = ''
        
        is_treat = treat_var and var_name == treat_var
        
        data.append({
            '变量': var_name,
            '系数': f"{coef:.4f}{sig}",
            '标准误': f"({se:.4f})",
            't 值': f"{t:.2f}",
            'P 值': f"{p:.4f}",
            '显著性': sig,
            '_is_treat': is_treat
        })
    
    df = pd.DataFrame(data)
    df_display = df[['变量', '系数', '标准误', 't 值', 'P 值']]
    return df_display, model.rsquared, model.rsquared_adj, len(model.resid)


# ============================================================
# Session State 初始化
# ============================================================
if 'df' not in st.session_state:
    st.session_state.df = None
if 'identified_vars' not in st.session_state:
    st.session_state.identified_vars = None
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = None
if 'base_model' not in st.session_state:
    st.session_state.base_model = None
if 'df_reg' not in st.session_state:
    st.session_state.df_reg = None

# ============================================================
# 主界面
# ============================================================

# 标题
st.markdown('<h1 class="main-header">📊 一键实证分析</h1>', unsafe_allow_html=True)
st.markdown('<p style="text-align: center; color: #666;">支持数据上传、变量自动识别、多元回归分析、稳健性检验、报告自动生成</p>', unsafe_allow_html=True)

# ============================================================
# 侧边栏 - 数据上传与变量选择
# ============================================================
with st.sidebar:
    st.markdown("### 📁 数据上传")
    
    uploaded_file = st.file_uploader(
        "上传数据文件",
        type=['dta', 'xlsx', 'xls', 'csv', 'sav'],
        help="支持 .dta (Stata)、.xlsx/.xls (Excel)、.csv、.sav (SPSS) 格式"
    )
    
    if uploaded_file is not None:
        # 获取文件类型
        file_type = os.path.splitext(uploaded_file.name)[1].lower()
        
        # 保存到临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_type) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        try:
            df = read_data(tmp_path, file_type)
            st.session_state.df = df
            
            # 自动识别变量
            identified = auto_identify_variables(df)
            st.session_state.identified_vars = identified
            
            st.success(f"✅ 数据加载成功！\n- 样本量: {len(df):,}\n- 变量数: {len(df.columns)}")
            
            # 显示数据预览
            with st.expander("📋 数据预览", expanded=False):
                st.dataframe(df.head(10), use_container_width=True)
            
        except Exception as e:
            st.error(f"❌ 文件读取失败：{str(e)}")
        
        finally:
            os.unlink(tmp_path)
    
    # 变量选择
    if st.session_state.df is not None:
        st.markdown("---")
        st.markdown("### 🎯 变量选择")
        
        df = st.session_state.df
        numeric_vars = df.select_dtypes(include=[np.number]).columns.tolist()
        all_vars = df.columns.tolist()
        
        # 自动识别结果
        identified = st.session_state.identified_vars
        
        # 因变量
        default_dep = identified['dep_var'] if identified['dep_var'] else numeric_vars[0] if numeric_vars else None
        dep_var = st.selectbox(
            "因变量 (Y)",
            options=numeric_vars,
            index=numeric_vars.index(default_dep) if default_dep in numeric_vars else 0,
            help="选择回归的因变量"
        )
        
        # 核心自变量
        default_treat = identified['treat_var'] if identified['treat_var'] else numeric_vars[1] if len(numeric_vars) > 1 else numeric_vars[0]
        treat_var = st.selectbox(
            "核心自变量 (X)",
            options=numeric_vars,
            index=numeric_vars.index(default_treat) if default_treat in numeric_vars else 0,
            help="选择核心解释变量"
        )
        
        # 控制变量
        default_controls = identified['control_vars'][:8] if identified['control_vars'] else []
        available_controls = [v for v in numeric_vars if v not in [dep_var, treat_var]]
        default_control_indices = [i for i, v in enumerate(available_controls) if v in default_controls]
        
        control_vars = st.multiselect(
            "控制变量",
            options=available_controls,
            default=[available_controls[i] for i in default_control_indices if i < len(available_controls)],
            help="选择控制变量（可多选）"
        )
        
        # 固定效应变量
        fe_vars = st.multiselect(
            "固定效应变量",
            options=all_vars,
            default=[v for v in [identified.get('region_var'), identified.get('time_var')] if v and v in all_vars],
            help="选择固定效应变量（如地区、年份）"
        )
        
        # 异质性分组变量
        group_var = st.selectbox(
            "异质性分组变量",
            options=['无'] + all_vars,
            index=0,
            help="选择分组变量进行异质性分析"
        )
        # 异质性分组变量
        group_var = st.selectbox(
            "异质性分组变量",
            options=['无'] + all_vars,
            index=0,
            help="选择分组变量进行异质性分析"
        )
        
        # 显著性优化选项
        st.markdown("---")
        st.markdown("### ✨ 显著性优化")
        enable_optimization = st.checkbox(
            "启用显著性优化",
            value=False,
            help="自动尝试多种方法优化回归结果"
        )
        
        if enable_optimization and OPTIMIZATION_AVAILABLE:
            target_p = st.selectbox(
                "目标显著性水平",
                options=[0.01, 0.05, 0.1],
                index=1,
                help="优化目标：P 值小于该水平"
            )
        elif enable_optimization and not OPTIMIZATION_AVAILABLE:
            st.warning("优化模块未找到，请确保 optimization_analysis.py 存在")
        else:
            target_p = 0.05
        # 分析按钮
        st.markdown("---")
        run_analysis = st.button("🚀 开始分析", type="primary", use_container_width=True)

# ============================================================
# 主内容区
# ============================================================

if st.session_state.df is None:
    # 欢迎页面
    st.markdown("""
    <div style="text-align: center; padding: 3rem;">
        <h2>欢迎使用一键实证分析工具</h2>
        <p style="color: #666; font-size: 1.1rem;">请在左侧上传数据文件开始分析</p>
    </div>
    """, unsafe_allow_html=True)
    
    # 功能介绍
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div class="info-box">
            <h4>📊 数据支持</h4>
            <ul>
                <li>Stata (.dta)</li>
                <li>Excel (.xlsx/.xls)</li>
                <li>CSV (.csv)</li>
                <li>SPSS (.sav)</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="info-box">
            <h4>🔬 分析功能</h4>
            <ul>
                <li>描述性统计</li>
                <li>基准回归</li>
                <li>异质性分析</li>
                <li>稳健性检验</li>
                <li>固定效应模型</li>
                <li>诊断检验</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="info-box">
            <h4>📥 输出结果</h4>
            <ul>
                <li>交互式图表</li>
                <li>回归结果表格</li>
                <li>Word 分析报告</li>
                <li>可复现代码</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)

else:
    # 运行分析
    if run_analysis or st.session_state.analysis_results is not None:
        
        if run_analysis:
            with st.spinner("正在进行实证分析..."):
                df = st.session_state.df
                results = {}
                
                # 1. 描述性统计
                numeric_vars = df.select_dtypes(include=[np.number]).columns.tolist()
                desc_stats = descriptive_statistics(df, numeric_vars[:15])
                results['descriptive'] = desc_stats
                results['data_info'] = {'n_obs': len(df), 'n_vars': len(df.columns)}
                
                # 2. 基准回归
                model, df_reg = baseline_regression(df, dep_var, treat_var, control_vars)
                results['baseline'] = {'model': model, 'n': len(df_reg)}
                st.session_state.base_model = model
                st.session_state.df_reg = df_reg
                
                # 3. 异质性分析
                if group_var != '无' and group_var in df.columns:
                    het_df = heterogeneity_analysis(df, dep_var, treat_var, control_vars, group_var)
                    results['heterogeneity'] = het_df
                else:
                    results['heterogeneity'] = None
                
                # 4. 稳健性检验
                robust_df = robustness_checks(df, dep_var, treat_var, control_vars)
                results['robustness'] = robust_df
                
                # 5. 固定效应模型
                fe_results = {}
                if fe_vars:
                    fe_model, _ = fixed_effects_regression(df, dep_var, treat_var, control_vars, fe_vars)
                    fe_results['model'] = fe_model
                    fe_results['fe_vars'] = fe_vars
                results['fixed_effects'] = fe_results if fe_vars else None
                
                # 6. 诊断检验
                diagnostics = {}
                try:
                    vif_vars = [treat_var] + control_vars[:5]
                    vif_result = calculate_vif(df_reg, vif_vars)
                    diagnostics['vif'] = vif_result
                except:
                    diagnostics['vif'] = None
                
                bp_result = breusch_pagan_test(model)
                diagnostics['bp_test'] = bp_result
                
                dw_result = durbin_watson_test(model)
                diagnostics['dw_test'] = dw_result
                
                results['diagnostics'] = diagnostics
                
                st.session_state.analysis_results = results
                st.session_state.dep_var = dep_var
                st.session_state.treat_var = treat_var
                st.session_state.control_vars = control_vars
                st.session_state.group_var = group_var
                st.session_state.group_var = group_var
                
                # 7. 显著性优化分析
                if enable_optimization and OPTIMIZATION_AVAILABLE:
                    with st.spinner("正在进行显著性优化分析（这可能需要几分钟）..."):
                        try:
                            opt_results, all_opt_results = significance_optimization(
                                df, dep_var, treat_var, control_vars,
                                target_p=target_p,
                                output_dir=tempfile.gettempdir()
                            )
                            st.session_state.optimization_results = opt_results
                            st.session_state.all_optimization_results = all_opt_results
                            st.success("✅ 显著性优化完成！")
                        except Exception as e:
                            st.error(f"优化分析失败：{str(e)}")
                            st.session_state.optimization_results = None
                
                st.success("✅ 分析完成！")
                st.success("✅ 分析完成！")
        
        # 显示结果
        results = st.session_state.analysis_results
        
        # 创建标签页
        # 创建标签页
        if st.session_state.get('optimization_results') is not None:
            tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
                "📈 基准回归", 
                "🔬 异质性分析", 
                "🛡️ 稳健性检验",
                "📐 固定效应",
                "🔍 诊断检验",
                "✨ 显著性优化",
                "📥 报告下载"
            ])
        else:
            tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
                "📊 描述性统计", 
                "📈 基准回归", 
                "🔬 异质性分析", 
                "🛡️ 稳健性检验",
                "📐 固定效应",
                "🔍 诊断检验",
                "📥 报告下载"
            ])
        
        # Tab 1: 描述性统计
        with tab1:
            st.markdown("### 描述性统计")
            
            desc = results['descriptive']
            
            # 格式化显示
            desc_display = desc.T.round(4)
            desc_display.columns = ['样本量', '均值', '标准差', '最小值', '25%', '中位数', '75%', '最大值',
                                    '偏度', '峰度', '缺失值', '缺失率(%)']
            
            st.dataframe(desc_display, use_container_width=True)
            
            # 数据可视化
            st.markdown("#### 变量分布可视化")
            
            numeric_cols = st.session_state.df.select_dtypes(include=[np.number]).columns.tolist()[:6]
            
            fig = make_subplots(rows=2, cols=3, subplot_titles=numeric_cols[:6])
            
            for i, col in enumerate(numeric_cols[:6]):
                row = i // 3 + 1
                col_idx = i % 3 + 1
                fig.add_trace(
                    go.Histogram(x=st.session_state.df[col].dropna(), name=col, showlegend=False),
                    row=row, col=col_idx
                )
            
            fig.update_layout(height=500, showlegend=False)
            st.plotly_chart(fig, use_container_width=True)
        
        # Tab 2: 基准回归
        with tab2:
            st.markdown("### 基准回归结果")
            
            model = results['baseline']['model']
            n = results['baseline']['n']
            treat_var = st.session_state.treat_var
            
            # 回归表格
            reg_table, r2, r2_adj, n_obs = format_regression_table(model, treat_var)
            
            st.dataframe(reg_table, use_container_width=True, hide_index=True)
            
            # 模型统计量
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("样本量", f"{n:,}")
            with col2:
                st.metric("R²", f"{r2:.4f}")
            with col3:
                st.metric("调整后 R²", f"{r2_adj:.4f}")
            with col4:
                st.metric("F 统计量", f"{model.fvalue:.2f}")
            
            st.markdown("""
            <p style="color: #666; font-size: 0.9rem;">
            注：*** p&lt;0.01, ** p&lt;0.05, * p&lt;0.1；使用稳健标准误 (HC3)
            </p>
            """, unsafe_allow_html=True)
            
            # 系数图
            st.markdown("#### 系数可视化")
            
            coef_data = []
            for var in model.params.index:
                if var != 'Intercept':
                    coef_data.append({
                        '变量': var,
                        '系数': model.params[var],
                        'se': model.bse[var],
                        'p': model.pvalues[var]
                    })
            
            coef_df = pd.DataFrame(coef_data)
            
            fig = go.Figure()
            
            colors = ['#E53935' if p < 0.05 else '#42A5F5' for p in coef_df['p']]
            
            fig.add_trace(go.Bar(
                x=coef_df['变量'],
                y=coef_df['系数'],
                error_y=dict(type='data', array=coef_df['se'] * 1.96),
                marker_color=colors
            ))
            
            fig.update_layout(
                title="回归系数 (95% 置信区间)",
                xaxis_title="变量",
                yaxis_title="系数",
                showlegend=False,
                height=400
            )
            
            st.plotly_chart(fig, use_container_width=True)
        
        # Tab 3: 异质性分析
        with tab3:
            st.markdown("### 异质性分析")
            
            if results['heterogeneity'] is not None and len(results['heterogeneity']) > 0:
                het_df = results['heterogeneity']
                
                st.dataframe(het_df.round(4), use_container_width=True, hide_index=True)
                
                # 可视化
                fig = go.Figure()
                
                fig.add_trace(go.Bar(
                    x=het_df['group'],
                    y=het_df['coef'],
                    error_y=dict(type='data', array=het_df['se'] * 1.96),
                    marker_color=['#E53935' if p < 0.05 else '#42A5F5' for p in het_df['pvalue']]
                ))
                
                fig.update_layout(
                    title=f"分组回归系数对比",
                    xaxis_title="组别",
                    yaxis_title="系数",
                    showlegend=False,
                    height=400
                )
                
                st.plotly_chart(fig, use_container_width=True)
                
                st.markdown("""
                <p style="color: #666; font-size: 0.9rem;">
                注：红色表示 p&lt;0.05，蓝色表示 p≥0.05
                </p>
                """, unsafe_allow_html=True)
            else:
                st.info("请在左侧选择分组变量进行异质性分析")
        
        # Tab 4: 稳健性检验
        with tab4:
            st.markdown("### 稳健性检验")
            
            robust_df = results['robustness']
            
            st.dataframe(robust_df.round(4), use_container_width=True, hide_index=True)
            
            # 可视化对比
            fig = go.Figure()
            
            fig.add_trace(go.Bar(
                x=robust_df['method'],
                y=robust_df['coef'],
                error_y=dict(type='data', array=robust_df['se'] * 1.96),
                marker_color=['#66BB6A' if p < 0.05 else '#BDBDBD' for p in robust_df['pvalue']]
            ))
            
            fig.update_layout(
                title="稳健性检验结果对比",
                xaxis_title="检验方法",
                yaxis_title="系数",
                showlegend=False,
                height=400,
                xaxis_tickangle=-45
            )
            
            st.plotly_chart(fig, use_container_width=True)
        
        # Tab 5: 固定效应模型
        with tab5:
            st.markdown("### 固定效应模型")
            
            if results['fixed_effects'] is not None:
                fe_model = results['fixed_effects']['model']
                fe_vars = results['fixed_effects']['fe_vars']
                
                st.markdown(f"**固定效应变量**: {', '.join(fe_vars)}")
                
                fe_table, fe_r2, fe_r2_adj, fe_n = format_regression_table(fe_model, st.session_state.treat_var)
                
                st.dataframe(fe_table, use_container_width=True, hide_index=True)
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("样本量", f"{fe_n:,}")
                with col2:
                    st.metric("R²", f"{fe_r2:.4f}")
                with col3:
                    st.metric("调整后 R²", f"{fe_r2_adj:.4f}")
            else:
                st.info("请在左侧选择固定效应变量")
        
        # Tab 6: 诊断检验
        with tab6:
            st.markdown("### 诊断检验")
            
            diag = results['diagnostics']
            
            col1, col2 = st.columns(2)
            
            # VIF 检验
            with col1:
                st.markdown("#### 多重共线性检验 (VIF)")
                
                if diag['vif'] is not None:
                    st.dataframe(diag['vif'], use_container_width=True, hide_index=True)
                    
                    # VIF 可视化
                    fig = go.Figure()
                    fig.add_trace(go.Bar(
                        x=diag['vif']['变量'],
                        y=diag['vif']['VIF'],
                        marker_color=['#E53935' if v >= 10 else ('#FB8C00' if v >= 5 else '#66BB6A') 
                                     for v in diag['vif']['VIF']]
                    ))
                    fig.add_hline(y=10, line_dash="dash", line_color="red", 
                                 annotation_text="高共线性阈值 (VIF=10)")
                    fig.add_hline(y=5, line_dash="dash", line_color="orange",
                                 annotation_text="中等共线性阈值 (VIF=5)")
                    fig.update_layout(title="VIF 值分布", height=350)
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.warning("VIF 计算失败，可能存在完全共线性")
            
            # BP 和 DW 检验
            with col2:
                st.markdown("#### 异方差检验 (Breusch-Pagan)")
                bp = diag['bp_test']
                
                bp_col1, bp_col2 = st.columns(2)
                with bp_col1:
                    st.metric("LM 统计量", f"{bp['lm_statistic']:.4f}")
                with bp_col2:
                    st.metric("P 值", f"{bp['lm_pvalue']:.4f}")
                
                if bp['lm_pvalue'] < 0.05:
                    st.warning(f"⚠️ 结论：{bp['conclusion']}，建议使用稳健标准误")
                else:
                    st.success(f"✅ 结论：{bp['conclusion']}")
                
                st.markdown("---")
                st.markdown("#### 自相关检验 (Durbin-Watson)")
                dw = diag['dw_test']
                
                st.metric("DW 统计量", f"{dw['dw_statistic']:.4f}")
                
                if dw['dw_statistic'] < 1.5 or dw['dw_statistic'] > 2.5:
                    st.warning(f"⚠️ 结论：{dw['conclusion']}")
                else:
                    st.success(f"✅ 结论：{dw['conclusion']}")

        
        # Tab 7: 显著性优化
        if st.session_state.get('optimization_results') is not None:
            with tab7:
                st.markdown("### ✨ 显著性优化分析")
                
                st.info("""
                **显著性优化说明**：本工具尝试多种方法（数据层面、模型层面、样本层面等）优化回归结果
                
                ⚠️ **学术诚信提醒**：本功能仅供学术研究和稳健性检验使用，请在学术诚信的前提下使用
                """)
                
                opt_results = st.session_state.optimization_results
                all_opt_results = st.session_state.all_optimization_results
                
                # 优化结果对比表格
                st.markdown("#### 优化结果对比")
                
                optimization_log = opt_results.get('optimization_log', [])
                if optimization_log:
                    opt_df = pd.DataFrame(optimization_log)
                    
                    # 添加显著性标记
                    def add_sig(p):
                        if p < 0.01:
                            return '***'
                        elif p < 0.05:
                            return '**'
                        elif p < 0.1:
                            return '*'
                        else:
                            return ''
                    
                    opt_df['显著性'] = opt_df['pvalue'].apply(add_sig)
                    
                    # 格式化显示
                    display_df = opt_df.copy()
                    display_df['系数'] = display_df['coef'].apply(lambda x: f"{x:.4f}")
                    display_df['标准误'] = display_df['se'].apply(lambda x: f"({x:.4f})")
                    display_df['P 值'] = display_df['pvalue'].apply(lambda x: f"{x:.4f}")
                    
                    display_cols = ['method', 'coef', 'se', 'pvalue', '显著性', 'n']
                    st.dataframe(display_df[display_cols], use_container_width=True, hide_index=True)
                    
                    st.markdown("注：*** p<0.01, ** p<0.05, * p<0.1")
                    
                    # 最优模型
                    if opt_results.get('best_model'):
                        st.markdown("#### 🏆 最优模型")
                        best = opt_results['best_model']
                        best_sig = add_sig(best['pvalue'])
                        
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("优化方法", best['method'])
                        with col2:
                            st.metric("系数", f"{best['coef']:.4f}{best_sig}")
                        with col3:
                            st.metric("P 值", f"{best['pvalue']:.4f}")
                    
                    # 诊断检验
                    diagnostics = opt_results.get('diagnostics', {})
                    if diagnostics:
                        st.markdown("#### 诊断检验")
                        
                        diag_col1, diag_col2 = st.columns(2)
                        
                        with diag_col1:
                            if 'vif' in diagnostics and diagnostics['vif'] is not None:
                                st.markdown("**VIF 多重共线性检验**")
                                st.dataframe(diagnostics['vif'], use_container_width=True, hide_index=True)
                        
                        with diag_col2:
                            if 'bp_test' in diagnostics:
                                bp = diagnostics['bp_test']
                                st.markdown("**BP 异方差检验**")
                                st.write(f"LM 统计量：{bp['lm_statistic']:.4f}, P 值：{bp['lm_pvalue']:.4f}")
                                st.write(f"结论：{bp['conclusion']}")
                            
                            if 'dw_test' in diagnostics:
                                dw = diagnostics['dw_test']
                                st.markdown("**DW 自相关检验**")
                                st.write(f"DW 统计量：{dw['dw_statistic']:.4f}")
                                st.write(f"结论：{dw['conclusion']}")
                    
                    # 优化方法说明
                    with st.expander("📖 优化方法说明"):
                        st.markdown("""
                        **数据层面优化**：
                        - Winsorize 缩尾（1%/5%）：处理极端值
                        - 对数转换：改善变量分布
                        
                        **模型层面优化**：
                        - 不同稳健标准误（HC1/HC2/HC3）：处理异方差
                        - 固定效应模型：控制个体/时间异质性
                        - 聚类标准误：处理组内相关
                        
                        **样本层面优化**：
                        - 剔除极端值：排除异常观测
                        
                        **诊断检验**：
                        - VIF：多重共线性诊断
                        - BP 检验：异方差检验
                        - DW 检验：自相关检验
                        """)
                else:
                    st.warning("未找到优化结果")
        
        # Tab 8: 报告下载（如果有优化结果）
        if st.session_state.get('optimization_results') is not None:
            with tab8:
                st.markdown("### 📥 分析报告下载")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("#### Word 分析报告")
                    st.markdown("包含完整的描述性统计、回归结果、稳健性检验等内容")
                    
                    # 生成报告
                    # 生成报告
                    if st.button("📄 生成 Word 报告", type="primary", key="word_report_btn"):
                        with st.spinner("正在生成报告..."):
                            word_buffer = generate_word_report(results)
                            
                            st.download_button(
                                label="📥 下载 Word 报告",
                                data=word_buffer,
                                file_name="实证分析报告.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    
                    # 显著性优化报告
                    st.markdown("---")
                    st.markdown("#### ✨ 显著性优化报告")
                    st.markdown("包含所有优化方法对比和诊断检验")
                    
                    if st.button("📄 生成显著性优化报告", type="primary", key="opt_report_btn"):
                        with st.spinner("正在生成优化报告..."):
                            try:
                                from optimization_analysis import generate_optimization_report
                                import tempfile
                                
                                opt_results = st.session_state.optimization_results
                                all_opt_results = st.session_state.all_optimization_results
                                
                                # 生成优化报告
                                report_path = os.path.join(tempfile.gettempdir(), '显著性优化报告.docx')
                                generate_optimization_report(opt_results, all_opt_results, report_path)
                                
                                # 读取文件供下载
                                with open(report_path, 'rb') as f:
                                    report_data = f.read()
                                
                                st.download_button(
                                    label="📥 下载显著性优化报告",
                                    data=report_data,
                                    file_name="显著性优化报告.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                            except Exception as e:
                                st.error(f"生成优化报告失败：{str(e)}")
                
                with col2:   col2:
                    st.markdown("#### 数据下载")   st.markdown   减价("#### Data Download")
                    
                    # 描述性统计
                    desc_csv = results['descriptive'].T.to_csv().encode('utf-8-sig')desc_csv = results['descriptive'   “描述性”].T.to_csv().encode   编码('utf-8-sig'   “utf-8-sig”)  # 将描述性统计结果转为 CSV 格式并编码为 UTF-8 带 BOM 格式
                    st.download_button(
                        label="📊 下载描述性统计 (CSV)",label=" 📊 Download Descriptive Statistics (CSV) ,
                        data=desc_csv,
                        file_name="描述性统计.csv",   file_name=" "Descriptive Statistics.csv" ,
                        mime="text/csv",   mime="text/csv",
                        key="desc_stats_download"
                    )

                    
                    # 回归结果
                    if results['baseline']:   mime="text/csv",   如果结果中存在“基准线”：
                        model = results['baseline']['model']
                        reg_summary = model.summary().as_csv()reg_summary = 模型总结().以 CSV 格式输出()
                        st.download_button(
                            label=   label=" 📈 Download Regression Results (CSV) ,"📈 下载回归结果 (CSV)",
                            data=reg_summary.encode   编码('utf-8-sig'   “utf-8-sig”),
                            file_name="回归结果.csv",
                            mime=   mime="text/csv","text/csv"   "text/csv",
                            key="reg_results_download"
                        )

        else:
            with   与 tab7:
                st.markdown   减价("### 📥 分析报告下载")
                
                col1, col2 = st.columns(2)
                
                with   与 col1:
                    st.markdown   减价("#### Word 分析报告")
                    st.markdown   减价("包含完整的描述性统计、回归结果、稳健性检验等内容")
                    
                    # 生成报告
                    # 生成报告
                    if st.button("📄 生成 Word 报告", type="primary", key="word_report_btn_no_opt"):if st.button("📄 Generate Word Report", type="primary", key="word_report_btn_no_opt"):
                        with st.spinner("正在生成报告..."):with st.spinner("Generating report...") "):
                            word_buffer = generate_word_report   生成单词报告(results)
                            
                            st.download_button(
                                label="📥 下载 Word 报告",   label=" Download the Word report. ,
                                data=word_buffer,   = word_buffer日期,
                                file_name="实证分析报告.docx",   file_name=" Empirical Analysis Report.docx ,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                
                with col2:   col2:
                    st.markdown   减价("#### 数据下载")
                    
                    # 描述性统计
                    desc_csv = results['descriptive'].T.to_csv().encode('utf-8-sig')desc_csv = results['descriptive'].T.to_csv().encode('utf-8-sig')  # 将描述性统计结果转为 CSV 格式并编码为 UTF-8 带 BOM 格式
                    st.download_button(
                        label="📊 下载描述性统计 (CSV)",label=" 📊 Download Descriptive Statistics (CSV) ,
                        data=desc_csv,
                        file_name="描述性统计.csv",   file_name=" "Descriptive Statistics.csv" ,
                        mime="text/csv",
                        key="desc_stats_download_no_opt"
                    )
                    
                    # 回归结果
                    if results['baseline']:   如果结果中存在“基准线”：
                        model = results['baseline']['model']
                        reg_summary = model.summary().as_csv()reg_summary = 模型总结().以 CSV 格式输出()
                        st.download_button(
                            label="📈 下载回归结果 (CSV)",   label=" 📈 Download Regression Results (CSV) ,
                            data=reg_summary.encode('utf-8-sig'),
                            file_name="回归结果.csv",
                            mime="text/csv",
                            key="reg_results_download_no_opt"
                        )

# ============================================================
# 页脚
# ============================================================
st.markdown   减价("---")
st.markdown   减价("""   st.markdown   """("""   """
<div style="text-align: center; color: #999; font-size: 0.85rem;">
    <p>一键实证分析工具 | 基于 Streamlit 构建 | 支持多种数据格式和分析方法</p>
</div>   < / div>
""", unsafe_allow_html=True)""", unsafe_allow_html=   真正的True   """)"""""", unsafe_allow_html=True), unsafe_allow_html=True)
